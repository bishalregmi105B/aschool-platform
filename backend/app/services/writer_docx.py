"""Writer v2 — server-side DOCX export (fallback for client-side export).

Mirrors frontend/lib/writer/exportDocx.ts: maps the TipTap ProseMirror JSON
plus WriterSettings onto a real .docx via python-docx. Accepts the payload
{doc, settings, name} posted by the frontend and returns the raw bytes.
"""

import base64
import io
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# px @96dpi → mm
MM_PER_PX = 25.4 / 96.0

PAGE_SIZES_MM = {
    "A4": (210.0, 297.0),
    "A5": (148.0, 210.0),
    "Letter": (215.9, 279.4),
    "Legal": (215.9, 355.6),
}

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_HEADING_COLORS = {
    1: "1E3A5F",
    2: "2C5282",
    3: "2D3748",
    4: "4A5568",
}


def _hex(color):
    """'#rrggbb' → 'RRGGBB' (or None)."""
    if not color:
        return None
    m = re.match(r"^#?([0-9a-fA-F]{6})$", str(color).strip())
    return m.group(1).upper() if m else None


def _set_paragraph_shading(paragraph, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _hex(color) or "FFFFFF")
    paragraph._p.get_or_add_pPr().append(shd)


def _set_paragraph_borders(paragraph, preset):
    """Preset borders — same mapping as the frontend writerParagraphFormat."""
    if not preset or preset == "none":
        return
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    sides = []
    if preset in ("top", "topbottom", "box"):
        sides.append("top")
    if preset in ("bottom", "topbottom", "box"):
        sides.append("bottom")
    if preset == "box":
        sides.extend(["left", "right"])
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "64748B")
        pBdr.append(el)
    pPr.append(pBdr)


def _set_columns(section, settings):
    """Multi-column section via w:cols (python-docx has no high-level API)."""
    count = int(settings.get("columns") or 1)
    if count <= 1:
        return
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    spacing_px = float(settings.get("columnSpacing") or 48)
    cols.set(qn("w:space"), str(int(round(spacing_px * MM_PER_PX * 56.7))))  # mm→twip
    if settings.get("columnDivider"):
        cols.set(qn("w:sep"), "1")


def _set_line_numbers(section):
    sect_pr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    sect_pr.append(ln)


def _add_page_field(paragraph):
    """Append a PAGE / NUMPAGES field pair ('Page X of Y')."""
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)

    fld_page = OxmlElement("w:fldSimple")
    fld_page.set(qn("w:instr"), " PAGE ")
    paragraph._p.append(fld_page)

    run2 = paragraph.add_run(" of ")
    run2.font.size = Pt(9)

    fld_pages = OxmlElement("w:fldSimple")
    fld_pages.set(qn("w:instr"), " NUMPAGES ")
    paragraph._p.append(fld_pages)


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_image(paragraph, src, width_px=None, height_px=None, max_width_mm=160):
    """Inline image from a data URL (PNG/JPEG/GIF/BMP).

    Honors editor-stored dimensions (px @96dpi → mm) when present, capped at
    ``max_width_mm`` so a full-width photo never overflows the page.
    """
    m = re.match(r"^data:image/(png|jpe?g|gif|bmp);base64,(.+)$", src or "", re.S | re.I)
    if not m:
        return
    ext = "jpg" if m.group(1).lower() in ("jpeg", "jpg") else m.group(1).lower()
    data = base64.b64decode(m.group(2))
    buf = io.BytesIO(data)
    buf.seek(0)
    try:
        kwargs = {}
        if width_px:
            w_mm = float(width_px) * MM_PER_PX
            if height_px:
                h_mm = float(height_px) * MM_PER_PX
                if w_mm > max_width_mm:  # keep aspect when clamping
                    h_mm = h_mm * max_width_mm / w_mm
                    w_mm = max_width_mm
                kwargs = {"width": Mm(w_mm), "height": Mm(h_mm)}
            else:
                kwargs = {"width": Mm(min(w_mm, max_width_mm))}
        paragraph.add_run().add_picture(buf, **kwargs)
    except Exception:  # corrupt image data — skip rather than fail the export
        pass


def _frame_paragraph(doc, attrs):
    """Floating box / wordart → Word text frame via w:framePr (best effort)."""
    text = str(attrs.get("text") or "")
    kind = str(attrs.get("kind") or "textbox")
    p = doc.add_paragraph()
    frame = p._p.get_or_add_pPr()
    fp = OxmlElement("w:framePr")
    fp.set(qn("w:x"), str(int(round(float(attrs.get("x") or 0) * MM_PER_PX * 56.7))))
    fp.set(qn("w:y"), str(int(round(float(attrs.get("y") or 0) * MM_PER_PX * 56.7))))
    fp.set(qn("w:w"), str(int(round(float(attrs.get("w") or 240) * MM_PER_PX * 56.7))))
    fp.set(qn("w:h"), str(int(round(float(attrs.get("h") or 100) * MM_PER_PX * 56.7))))
    fp.set(qn("w:hRule"), "atLeast")
    fp.set(qn("w:xAlign"), "left")
    fp.set(qn("w:yAlign"), "inline")
    frame.insert(0, fp)

    if kind == "rect" or kind == "ellipse" or kind == "arrow" or kind == "star":
        if attrs.get("fill"):
            _set_paragraph_shading(p, attrs.get("fill"))
        return p
    if kind == "wordart":
        run = p.add_run(text or "WordArt")
        run.bold = True
        run.font.size = Pt(max(float(attrs.get("fontSize") or 36), 24))
        if attrs.get("color"):
            run.font.color.rgb = RGBColor.from_string(_hex(attrs.get("color")) or "1D4ED8")
        return p

    run = p.add_run(text)
    run.font.size = Pt(float(attrs.get("fontSize") or 12))
    if attrs.get("color"):
        run.font.color.rgb = RGBColor.from_string(_hex(attrs.get("color")) or "0F172A")
    if attrs.get("border"):
        _set_paragraph_borders(p, "box")
    return p


def _apply_run_marks(run, marks, settings):
    bold = False
    italic = False
    underline = False
    strike = False
    sub = False
    sup = False
    color = None
    highlight = None
    font = None
    size = None
    for m in marks or []:
        mtype = m.get("type")
        attrs = m.get("attrs") or {}
        if mtype == "bold":
            bold = True
        elif mtype == "italic":
            italic = True
        elif mtype == "underline":
            underline = True
        elif mtype == "strike":
            strike = True
        elif mtype == "subscript":
            sub = True
        elif mtype == "superscript":
            sup = True
        elif mtype == "color":
            color = attrs.get("color")
        elif mtype == "highlight":
            highlight = attrs.get("color")
        elif mtype == "fontFamily":
            font = attrs.get("fontFamily")
        elif mtype == "fontSize":
            raw = attrs.get("fontSize")
            try:
                size = float(str(raw).replace("pt", ""))
            except (TypeError, ValueError):
                size = None
    run.bold = bold or None
    run.italic = italic or None
    if underline:
        run.underline = True
    if strike:
        rPr = run._r.get_or_add_rPr()
        rPr.append(OxmlElement("w:strike"))
    if sub:
        run.font.subscript = True
    if sup:
        run.font.superscript = True
    if color and _hex(color):
        run.font.color.rgb = RGBColor.from_string(_hex(color))
    if highlight and _hex(highlight):
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), _hex(highlight))
        rPr.append(shd)
    if font:
        run.font.name = font
    if size:
        run.font.size = Pt(size)


def _walk_inline(paragraph, node, settings):
    """Add runs for an inline node tree (text + marks + links + images)."""
    for child in node.get("content") or []:
        ctype = child.get("type")
        if ctype == "hardBreak":
            paragraph.add_run().add_break()
        elif ctype == "image":
            attrs = child.get("attrs") or {}
            _add_image(paragraph, str(attrs.get("src") or ""), attrs.get("width"), attrs.get("height"))
        elif child.get("text") is not None:
            marks = child.get("marks") or []
            link = next((m for m in marks if m.get("type") == "link"), None)
            if link:
                _add_hyperlink(paragraph, str((link.get("attrs") or {}).get("href") or "#"), child.get("text"))
                continue
            run = paragraph.add_run(child.get("text"))
            _apply_run_marks(run, marks, settings)


def _para_props(paragraph, attrs, settings):
    if attrs.get("textAlign") in _ALIGN:
        paragraph.alignment = _ALIGN[attrs["textAlign"]]
    fmt = paragraph.paragraph_format
    try:
        fmt.line_spacing = float(attrs.get("lineHeight") or 1.6)
    except (TypeError, ValueError):
        fmt.line_spacing = 1.6
    indent = attrs.get("indent")
    if indent:
        fmt.left_indent = Mm(float(indent) * MM_PER_PX)
    if attrs.get("borders"):
        _set_paragraph_borders(paragraph, attrs.get("borders"))
    if attrs.get("shading"):
        _set_paragraph_shading(paragraph, attrs.get("shading"))


def _render_blocks(doc, blocks, settings, depth=0):
    list_counter = {"n": 0}
    for node in blocks or []:
        ntype = node.get("type")
        attrs = node.get("attrs") or {}

        if ntype == "pageBreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        elif ntype == "heading":
            level = min(4, max(1, int(attrs.get("level") or 1)))
            p = doc.add_paragraph(style=f"Heading {level}")
            _para_props(p, attrs, settings)
            _walk_inline(p, node, settings)

        elif ntype == "paragraph":
            p = doc.add_paragraph()
            _para_props(p, attrs, settings)
            _walk_inline(p, node, settings)

        elif ntype == "blockquote":
            p = doc.add_paragraph()
            _para_props(p, attrs, settings)
            p.paragraph_format.left_indent = Mm(12 + float(attrs.get("indent") or 0) * MM_PER_PX)
            _walk_inline(p, node, settings)

        elif ntype == "codeBlock":
            code = "\n".join((c.get("text") or "") for c in node.get("content") or [])
            p = doc.add_paragraph()
            _set_paragraph_shading(p, "#F1F5F9")
            run = p.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(10)

        elif ntype == "horizontalRule":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), "94A3B8")
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif ntype == "image":
            attrs = node.get("attrs") or {}
            p = doc.add_paragraph()
            if attrs.get("textAlign") in _ALIGN:
                p.alignment = _ALIGN[attrs["textAlign"]]
            _add_image(p, str(attrs.get("src") or ""), attrs.get("width"), attrs.get("height"))

        elif ntype == "bulletList":
            for li in node.get("content") or []:
                list_counter["n"] += 1
                _render_list_item(doc, li, settings, "List Bullet", depth)

        elif ntype == "orderedList":
            for li in node.get("content") or []:
                _render_list_item(doc, li, settings, "List Number", depth)

        elif ntype == "table":
            _render_table(doc, node, settings)

        elif ntype == "floatingBox":
            _frame_paragraph(doc, node.get("attrs") or {})

        else:
            # unknown block — flatten its text
            p = doc.add_paragraph()
            _walk_inline(p, node, settings)


def _render_list_item(doc, li, settings, style, depth):
    style_name = style if depth == 0 else f"{style} {min(depth + 1, 3)}"
    for child in li.get("content") or []:
        ctype = child.get("type")
        if ctype == "paragraph":
            p = doc.add_paragraph()
            try:
                p.style = doc.styles[style_name]
            except KeyError:
                p.paragraph_format.left_indent = Mm(10 * (depth + 1))
            _para_props(p, child.get("attrs") or {}, settings)
            _walk_inline(p, child, settings)
        elif ctype in ("bulletList", "orderedList"):
            sub_style = "List Bullet" if ctype == "bulletList" else "List Number"
            for sub_li in child.get("content") or []:
                _render_list_item(doc, sub_li, settings, sub_style, depth + 1)
        elif ctype == "pageBreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _render_table(doc, node, settings):
    rows = [r for r in node.get("content") or [] if r.get("type") == "tableRow"]
    if not rows:
        return
    cols = max(len(r.get("content") or []) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row.get("content") or []):
            if ci >= cols:
                break
            target = table.cell(ri, ci)
            if cell.get("type") == "tableHeader":
                target.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            target.text = ""
            para = target.paragraphs[0]
            for block in cell.get("content") or []:
                if block.get("type") == "paragraph":
                    _para_props(para, block.get("attrs") or {}, settings)
                    _walk_inline(para, block, settings)
                else:
                    _walk_inline(para, block, settings)
            if cell.get("type") == "tableHeader":
                for run in para.runs:
                    run.bold = True


def writer_doc_to_bytes(doc, settings, name="document"):
    """Render the writer2 TipTap JSON + settings into .docx bytes."""
    settings = settings or {}
    result = Document()

    section = result.sections[0]
    size = PAGE_SIZES_MM.get(settings.get("pageSize") or "A4", PAGE_SIZES_MM["A4"])
    portrait = (settings.get("orientation") or "portrait") == "portrait"
    width_mm, height_mm = size if portrait else (size[1], size[0])
    section.page_width = Mm(width_mm)
    section.page_height = Mm(height_mm)
    if not portrait:
        section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Mm(float(settings.get("marginTop") or 96) * MM_PER_PX)
    section.right_margin = Mm(float(settings.get("marginRight") or 76) * MM_PER_PX)
    section.bottom_margin = Mm(float(settings.get("marginBottom") or 96) * MM_PER_PX)
    section.left_margin = Mm(float(settings.get("marginLeft") or 76) * MM_PER_PX)

    _set_columns(section, settings)
    if settings.get("lineNumbers"):
        _set_line_numbers(section)

    # base font
    normal = result.styles["Normal"]
    normal.font.name = settings.get("font") or "Calibri"
    normal.font.size = Pt(float(settings.get("fontSize") or 11))

    # header
    if settings.get("headerOn") and settings.get("headerText"):
        hp = section.header.paragraphs[0]
        hp.text = settings.get("headerText")
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # footer (+ page numbers)
    footer_children_used = bool(settings.get("footerOn") and settings.get("footerText"))
    fp = section.footer.paragraphs[0]
    fp.alignment = _ALIGN.get(
        str(settings.get("pageNumber") or "none").replace("bottom-", ""),
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    if footer_children_used:
        run = fp.add_run(settings.get("footerText"))
        run.font.size = Pt(9)
    if settings.get("pageNumber") not in (None, "none"):
        _add_page_field(fp)

    blocks = (doc or {}).get("content") or []
    if not blocks:
        result.add_paragraph("")
    else:
        _render_blocks(result, blocks, settings)

    buf = io.BytesIO()
    result.save(buf)
    buf.seek(0)
    return buf.read()
